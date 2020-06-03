#!C:\Program Files\Python38\python.exe

#Converting List to String
def listToString(s):  
    
    # initialize an empty string 
    str1 = ""  
    
    # traverse in the string   
    for ele in s:  
        str1 += ele   
    
    # return string   
    return str1 
from docx import Document
import comtypes.client
import pymysql as sql
import random
from datetime import datetime
print("\n\n\n")
t=0
#Creating of Secret Key based on timing to get unique Names
now = datetime.now()
year = now.strftime("%Y")
month = now.strftime("%m")
time = now.strftime("%H:%M:%S")
times=""
for i in range(0,7):
    if(time[i]!=':'):
        times=times+time[i]
key=year+month+times
#Connecting DB
db=sql.connect('localhost','root','','myminiproject',port=3307)
cursor=db.cursor()
#Creating Document
doc= Document()
doc.add_heading("                                        MID SEM QUESTION PAPER\n Marks:70                                                      TIME:180MINS",level=1)
names=[]#To store Chapter Secret key
one=[]#TO store number of question
two=[]
three=[]
five=[]
Qone=[]#To store the questions accordig to marks
Qtwo=[]
Qthree=[]
Qfive=[]
Lone=[]#To store number of questions available in each module
Ltwo=[]
Lthree=[]
Lfive=[]
#Storing the number chapters of blueprint using Max funtion
query="CREATE VIEW MAXID AS(SELECT MAX(SlNo) Maximum from blueprintforqp)"
#print(query)
cursor.execute(query)
#getting number of chapters by using Max function
query="Select SlNo from blueprintforqp,MAXID where SlNo=Maximum"
cursor.execute(query)
r = cursor.fetchall()
#Droping create view
query="DROP VIEW MAXID"
cursor.execute(query)
#COllecting Secret key and storing it in names
for row in r:
   num = row[0]
   for i in range(1,(num+1)):
     query='Select SK from blueprintforqp where SlNo='+str(i)
     cursor.execute(query)
     r = cursor.fetchall()
     for row in r:
         names.append(row[0])
     chapter=len(names)#Stores Number of chapters
   #Number of Questions asked in Blueprint 
   for i in range (0,chapter):     
        query="Select One,Two,Three,Five from blueprintforqp where SK="+str(names[i])
        #print(query)
        cursor.execute(query)
        r = cursor.fetchall()
        #Number of questions mentioned in blueprint according to marks
        for row in r:
            one.append(row[0])
            two.append(row[1])
            three.append(row[2])
            five.append(row[3])
        #Storing the Questions from DB in QOne
        query="Select Questions from `"+str(names[i])+"` where Marks=1"
        cursor.execute(query)
        r = cursor.fetchall()
        for row in r:
           Qone.append([row[0]] )
          #Storing number of question in Lone
        Lone.append(len(Qone))
        query="Select Questions from `"+str(names[i])+"` where Marks=2"
        cursor.execute(query)
        r = cursor.fetchall()
        for row in r:
           Qtwo.append([row[0]] )
        Ltwo.append(len(Qtwo))
        query="Select Questions from `"+str(names[i])+"` where Marks=3"
        cursor.execute(query)
        r = cursor.fetchall()
        for row in r:
           Qthree.append([row[0]] )
        Lthree.append(len(Qthree))
        query="Select Questions from `"+str(names[i])+"` where Marks=5"
        cursor.execute(query)
        r = cursor.fetchall()
        for row in r:
           Qfive.append([row[0]] )
        Lfive.append(len(Qfive))
Norepeat=[]
doc.add_heading("               PART-A",level=2)
doc.add_heading("   Answer the following questions in one more or sentence",level=2)
doc.add_heading("                                                                                    (10*1=10)",level=4)
#Selecting random question
try:    
    for i in range(0,chapter):
        if(one[i]>0):
            n=one[i]
            if (i==0):
                mod=Lone[i]
            else:
                mod=Lone[i]-Lone[i-1]
            Norepeat=[]
            #checks whether required questions are thier in DB or no
            if(mod>=n):
                for j in range(0,n):
                    t=1
                    #Avoid duplicate of question
                    while(t==1):
                        #Generate random number
                        ques1=random.randrange(100,999)
                        if(i==0):               
                            na=(ques1%mod)+1  
                        else:
                            na=(ques1%mod)+Lone[i-1]
                        if na in Norepeat:
                            continue
                        else:
                           t=2
                           Norepeat.append(na)
                    doc.add_paragraph(listToString(Qone[na]),style='List Number')
        else:
            print("Question paper cant be created")
            exit()
except:
    print(" ")
doc.add_heading("   Answer any five of the following in 3 to 5 sentences",level=2)
doc.add_heading("                                                                                          (2*5=10)",level=4)
Norepeat=[]
try:
    for i in range(0,chapter):
        if(mod>=n):
            n=two[i]
            if (i==0):
                mod=Ltwo[i]
            else:
                mod=Ltwo[i]-Ltwo[i-1]
            Norepeat=[]
            for j in range(0,n):
                t=1
                while(t==1):
                    ques1=random.randrange(100,999)
                    if(i==0):               
                        na=(ques1%mod)+1  
                    else:
                        na=(ques1%mod)+Ltwo[i-1]
                    if na in Norepeat:
                        continue
                    else:
                       t=2
                       Norepeat.append(na)
                doc.add_paragraph(listToString(Qtwo[na]),style='List Number')
        else:
            print("Question paper cant be created")
            exit() 
except:
    print(" ")
doc.add_heading("   PART-C",level=2)
doc.add_heading("   Answer any five of the following in 40 to 80 words",level=2)
doc.add_heading("                                                                                        (3*5=15)",level=4)
Norepeat=[]
try:
    for i in range(0,chapter):
        n=three[i]
        if (i==0):
            mod=Lthree[i]
        else:
            mod=Lthree[i]-Lthree[i-1]
        Norepeat=[]
        if(mod>=n):
            for j in range(0,n):
                t=1
                while(t==1):
                    ques1=random.randrange(100,999)
                    if(i==0):               
                        na=(ques1%mod)
                    else:
                        na=(ques1%mod)+Lthree[i-1]
                    if na in Norepeat:
                        continue
                    else:
                       t=2
                       Norepeat.append(na)
                doc.add_paragraph(listToString(Qthree[na]),style='List Number')
        else:
            print("Question paper cant be created")
            exit() 
except:
    print(" ")
doc.add_heading("   PART-C",level=2)
doc.add_heading("   Answer any seven of the following in 200 to 250 words",level=2)
doc.add_heading("                                                                                        (5*7=35)",level=4)
Norepeat=[]
try:
    for i in range(0,chapter):
        n=five[i]
        if (i==0):
            mod=Lfive[i]
        else:
            mod=Lfive[i]-Lfive[i-1]
        Norepeat=[]
        if(mod>=n):
            for j in range(0,n):
                t=1
                while(t==1):
                    ques1=random.randrange(100,999)
                    if(i==0):               
                        na=(ques1%mod)+1  
                    else:
                        na=(ques1%mod)+Lfive[i-1]
                    if na in Norepeat:
                        continue
                    else:
                       t=2
                       Norepeat.append(na)
                doc.add_paragraph(listToString(Qfive[na]),style='List Number')
        else:
            print("Question paper cant be created")
            exit()
except:
    print(" ")
doc.save("'" +key+"_qpaper'.docx")
in_file=r"C:\xampp\htdocs\QPG\'" +key+"_qpaper'.docx"
out=r"C:\xampp\htdocs\QPG\QuestionPaper\'" +key+"_qpaper'.pdf"
word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out, FileFormat=17)
doc.Close()
word.Quit()
import magic
mime = magic.Magic(mime=True)
mof=mime.from_file(out)
import mysql.connector
def convertToBinaryData(filename):
    # Convert digital data to binary format
    with open(filename, 'rb') as file:
        binaryData = file.read()
    return binaryData

def insertBLOB(name,ftype,File):
   # print("Inserting BLOB into QuestionBank table")
    try:
        connection = mysql.connector.connect(host='localhost',
                                             database='myminiproject',
                                             user='root',
                                             password='',
                                             port=3307)

        cursor = connection.cursor()
        sql_insert_blob_query = """ INSERT INTO questionpaper
                          (FileName,FileType,FileData) VALUES (%s,%s,%s)"""

        file = convertToBinaryData(File)

        # Convert data into tuple format
        insert_blob_tuple = (name,ftype,file)
        cursor.execute(sql_insert_blob_query, insert_blob_tuple)
        connection.commit()
       # print("File inserted successfully as a BLOB into QuestionBank table", result)

    except mysql.connector.Error as error:
        print("Failed inserting BLOB data into MySQL table {}".format(error))

    finally:
        if (connection.is_connected()):
            cursor.close()
            connection.close()
   

insertBLOB(""+key+"",mof,out)   
print("Question Paper is created")
print("Secret key is ",key)#!C:\Program Files\Python38\python.exe

