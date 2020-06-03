#!C:\Program Files\Python38\python.exe
from datetime import datetime
print("\n\n\n")
import docx
import comtypes.client
from docx import Document
import pymysql as sql
import pymysql.cursors
def write_file(data, filename):
    # Convert binary data to proper format and write it on Hard Disk
    with open(filename, 'wb') as file:
        file.write(data)
db=sql.connect('localhost','root','','myminiproject',port=3307)
cursor=db.cursor()
query="CREATE VIEW MAXSlNo AS(SELECT MAX(SlNo) Maximum from inputdocumentdata)"
cursor.execute(query)
query="Select FileData from inputdocumentdata,MAXSlNo where SlNo=Maximum"
cursor.execute(query)
r = cursor.fetchall()
query="DROP VIEW MAXSlNo"
cursor.execute(query)
for row in r:
    file = row[0]
    file = row[0]
    ###print(file)  
    write_file(file,"extracteddata.docx")    
doc=docx.Document('extracteddata.docx')
now = datetime.now()
t=0
year = now.strftime("%Y")
month = now.strftime("%m")

time = now.strftime("%H:%M:%S")

times=""
for i in range(0,7):
    if(time[i]!=':'):
        times=times+time[i]
key=year+month+times
#print("\n\n\n\nSecret key is "+key+"\n\n")
keys=key
key="`"+key+"`"
sql_1="CREATE TABLE "+key+" ( SlNo int(5) NOT NULL AUTO_INCREMENT, Questions LONGTEXT NOT NULL,  Marks int(2) NOT NULL,PRIMARY KEY (SlNo))"
cursor.execute(sql_1)
i=0
c=[]
doc= Document()
doc.add_heading('                                        Question Bank\n', level=1)
for paragraphs in doc.paragraphs:    
    c=(doc.paragraphs[i].text)
    i=i+1
document =docx.Document('extracteddata.docx')
num=[]
specialChar=[]
specialChar=[':',':-','-']
num=['0','1','2','3','4','5','6','7','8','9']
a=[]
heading=[]
count=0
Insertdata=[]
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            if run.text[0] in num:
                heading.append(run.text)
            else:
                headlen=len(run.text)
                for i in range(0,headlen):
                    kpa=len(run.text)
                    try:
                        if run.text[0].isupper() and run.text[i]==' ' and ((kpa)>i):
                            if(run.text[i+1].isupper()):
                                for j in range(i+1,headlen):
                                    if run.text[i+1].isupper() and run.text[j]==' ' and run.text[j+1].isupper():
                                        heading.append(run.text)
                    except:
                        excep=0
                if run.text[headlen-1] in specialChar:
                    heading.append(run.text)
                if run.text not in heading and run.text[0]!='F':
                        a.extend("_")
        else:
            a.extend(run.text)
###print(heading)
#counting number of sentences under particular heading
co=[]
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.bold and run.text in heading:
                co.append(count)
                count=0
        else:    
            b=run.text
            leng=len(run.text)
            for i in range(0,leng):
                if b[i]=='.' or b[i]==';':
                    if i<leng:
                        try:
                            if b[i+1]==' ':
                                count=count+1
                        except:
                            count=count+1
            
co.append(count)
#extracting italics from document
ques=[]
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.italic:
            b=len(run.text)
            if run.text[b-1]=='?' and b>=10:
                ques.append(run.text)
len1=len(ques)
two=[]
one=[]
for i in range(0,len1):
    leng1=len(ques[i])
    count=1
    for j in range(0,leng1-1):
        if ques[i][j]=="?":
            count=count+1
    if count>=2 or ques[i][0]=='C':
        two.append(ques[i])
    else:
        one.append(ques[i])
#to extract only bold lines
b=len(a)
onlybold=[]
c=1
m=0
#extracting only lines which has _
for i in range(0,b):
    if a[i]=='_':
        j=i
        while a[j]!='.':
            j=j+1
        k=i
        while a[k]!='.':
            k=k-1
        for n in range(k,j):
            if a[n]!='.':
                ###print(a[n],end='')
                onlybold.insert(m,a[n])
            else:
                ###print(".\n\n",c,")",end='')
                onlybold.insert(m,'.')
                c=c+1
            n=n+1
            m=m+1
m=m+1
onlybold.insert(m,'.')
#iserting ________ for _
for i in range(1,m):
    if onlybold[i]=='_':
        onlybold[i]="________"
###printing it
result=[]
aresString=""
##print("\nFill in the Blank\n")
doc.add_heading("\nFill in the Blank\n",level=2)
for i in range(1,m):
    aresString=""
    if onlybold[i]=='.':
        j=i-1
        while onlybold[j]!='.':
            j=j-1
        for d in range(j+1,i+1):
            aresString=aresString+onlybold[d]
        result.append(aresString)
result=list(dict.fromkeys(result))
reslen=len(result)
n=1
for i in range(0,reslen):
    lenres=len(result[i])
    if result[i][lenres-1]=='.':
        doc.add_paragraph(result[i],style='List Number')
        ##print(n,")",end='')
        ##print(result[i])
        n=n+1
#for one and five marks questions
#extrcting bold letters
boldlist=[]
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            boldlist.append(run.text)
#removal of numbers and special characters from all bold lettered words
num=[1,2,3,4,5,6,7,8,9,'.']
j=0
boldword=[]
string=""
for m in boldlist:
    k=len(m)
    for i in range(0,k):
        if(boldlist[j][i]=='1' or boldlist[j][i]=='.' or boldlist[j][i]==':' or boldlist[j][i]=='2' or boldlist[j][i]=='3' or boldlist[j][i]=='4' or boldlist[j][i]=='5' or boldlist[j][i]=='6' or boldlist[j][i]=='7' or boldlist[j][i]=='8' or boldlist[j][i]=='9' or boldlist[j][i]=='0'):
            continue;
        else:
            if(string==""):
                if(boldlist[j][i]==' '):
                    continue;
                else:
                    string=string+boldlist[j][i] 
            else:
                string=string+boldlist[j][i] 
    j=j+1
    if len(string)!=1:
        boldword.append(string)
    string=""
#removal of numbers and special characters from heading
j=0
head=[]
string=""
for el in heading:
    k=len(el)
    for i in range(0,k):
        if(heading[j][i]=='1' or heading[j][i]=='.' or heading[j][i]==':' or heading[j][i]=='2' or heading[j][i]=='3' or heading[j][i]=='4' or heading[j][i]=='5' or heading[j][i]=='6' or heading[j][i]=='7' or heading[j][i]=='8' or heading[j][i]=='9' or heading[j][i]=='0'):
            continue;
        else:
            if(string==""):
                if(heading[j][i]==' '):
                    continue;
                else:
                    string=string+heading[j][i] 
            else:
                string=string+heading[j][i] 
    j=j+1
    if len(string)!=1:
        head.append(string)
    string=""
#one mark question
##print(co)
##print("\nOne marks Questions\n")
marks="1"
doc.add_heading("\nOne marks Questions\n",level=2)
boldword=list(dict.fromkeys(boldword))
len1=len(boldword)
#one marrk question
n=1
for i in range(0,len1):
    if boldword[i]!= 'Figure ' and boldword[i] not in head:
        str="What is "+boldword[i]+" ?"
        doc.add_paragraph(str,style='List Number')
        ##print(n,") What is ",end='')
        ##print(boldword[i],end='')
        ##print("?")
        sql_2="INSERT into "+key+" (Questions,Marks)VALUES('"+str+"',"+marks+")"
        cursor.execute(sql_2)
        db.commit()
        n=n+1
#duplicates removal
head=list(dict.fromkeys(head))
lenght=len(head)
#Two marks question
n=1
##print("\nTwo marks Questions\n")
doc.add_heading("\nTwo marks Questions\n",level=2)
for i in range(0,lenght):
    if co[i+1]>=2 and co[i+1]<=3:
        str="Define "+head[i]+" ?"
       ## print(n,") Define ",end='')
       # document.add_paragraph(n)
        doc.add_paragraph(str,style='List Number')
        ##print(head[i],end='')
        ##print("?")
        n=n+1
k=int(len1/2)
marks="2"
for i in range(0,k):
    if boldword[i]!= 'Figure ' and boldword[i] not in head:
        str="What is "+boldword[i]+" and "+boldword[k+i]+" ?"
        doc.add_paragraph(str,style='List Number')
        ##print(n,") What is ",end='')
        ##print(boldword[i],end='')
        ##print("?")       
        n=n+1
        sql_2="INSERT into "+key+" (Questions,Marks)VALUES('"+str+"',"+marks+")"
        cursor.execute(sql_2)
        db.commit()
#Three marks question
n=1
marks="3"
doc.add_heading("\nThree marks Questions\n",level=2)
for i in range(0,lenght):
    if co[i+1]>=3 and co[i+1]<=6:
        ##print(n,") Explain ",end='')
        str="Explain "+head[i]+ "?"
        doc.add_paragraph(str,style='List Number')
        ##print(head[i],end='')
        ##print("?\n")
        sql_2="INSERT into "+key+" (Questions,Marks)VALUES('"+str+"',"+marks+")"
        cursor.execute(sql_2)
        db.commit()
        n=n+1
lenone=len(one)
lentwo=len(two)
if(lenone>lentwo):
    minlen=lentwo
else:
    minlen=lenone
for i in range(0,minlen):
    str= "i)"+one[i]+"                   (1+2)\nii)"+two[i]
    sql_2="INSERT into "+key+" (Questions,Marks)VALUES('"+str+"',"+marks+")"
    doc.add_paragraph(str,style='List Number')
    cursor.execute(sql_2)
    db.commit()
minlen=(lenone-lentwo)
n=minlen
for i in range((minlen+3),lenone,3):
    str= "i)"+one[n]+"                   (1+1+1)\nii)"+one[n+1]+"\niii)"+one[n+2]
    doc.add_paragraph(str,style='List Number')
    sql_2="INSERT into "+key+" (Questions,Marks)VALUES('"+str+"',"+marks+")"
    cursor.execute(sql_2)
    db.commit()
    n=n+3

#Five marks question

n=1
marks="5"
doc.add_heading("\nFive marks Questions\n",level=2)
##print("\nFive marks Questions\n")
for i in range(0,lenght):
    if co[i+1]>=6:
        ##print(n,") Explain ",end='')
        str="Explain "+head[i]+ "?"
        doc.add_paragraph(str,style='List Number')
        ##print(head[i],end='')
        ##print("?\n")
        n=n+1
        sql_2="INSERT into "+key+" (Questions,Marks)VALUES('"+str+"',"+marks+")"
        cursor.execute(sql_2)
        db.commit()
sql_2="DELETE FROM blueprintforqp"
cursor.execute(sql_2)
doc.save("'"+keys+"'_qbank.docx")
in_file=r"C:\xampp\htdocs\QPG\'"+keys+"'_qbank.docx"
out=r"C:\xampp\htdocs\QPG\QuestionBank\'"+keys+"'_qbank.pdf"
word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out, FileFormat=17)
doc.Close()
word.Quit()
import magic
mime = magic.Magic(mime=True)
mof=mime.from_file(out)
#print(mof)
import mysql.connector
def convertToBinaryData(filename):
    # Convert digital data to binary format
    with open(filename, 'rb') as file:
        binaryData = file.read()
    return binaryData

def insertBLOB(name,ftype,File):
 #   print("Inserting BLOB into QuestionBank table")
    try:
       
        sql_insert_blob_query = """ INSERT INTO questionbank
                          (FileName,FileType,FileData) VALUES (%s,%s,%s)"""

        file = convertToBinaryData(File)

        # Convert data into tuple format
        insert_blob_tuple = (name,ftype,file)
        result = cursor.execute(sql_insert_blob_query, insert_blob_tuple)
        db.commit()
    #    print("File inserted successfully as a BLOB into QuestionBank table", result)

    except mysql.connector.Error as error:
        t=1
        print("Failed inserting BLOB data into MySQL table {}".format(error))


insertBLOB(""+keys+"",mof,out)
if t==0:
    print("\n\nThe secret key is ",keys)
    print("please store it properly to use next time\n\n")